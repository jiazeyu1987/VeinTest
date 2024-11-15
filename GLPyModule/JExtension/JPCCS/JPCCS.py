import imp
import re
import os
from re import A
from tabnanny import check
from time import sleep
from datetime import datetime
import unittest
import logging
import vtk, qt, ctk, slicer
from slicer.ScriptedLoadableModule import *
from slicer.util import VTKObservationMixin
import slicer.util as util
import SlicerWizard.Utilities as su
import numpy as np
import SimpleITK as sitk
from ConfigTools.DetailInfo import HistoryDetailPanel

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


#
# JPCCS
#


class JNewPatient(qt.QDialog, ScriptedLoadableModuleWidget):
    user_id = 0
    patient_info = []
    gender_list = ["男", "女", "其他"]
    # 治疗位置
    treat_type_list = ["左侧", "右侧"]
    # 临床分类
    degree_list = ["C1", "C2", "C3", "C4", "C5", "C6"]

    def __init__(self, parent=None):
        super(JNewPatient, self).__init__()
        self.setWindowFlag(qt.Qt.FramelessWindowHint)
        self.setAttribute(qt.Qt.WA_TranslucentBackground)  # 使窗口透明

        # 设置对话框为全屏
        # self.showFullScreen()
        self.module_path = os.path.dirname(slicer.util.modulePath("JPCCS"))
        print(self.module_path)
        uiWidget = slicer.util.loadUI(self.module_path + '/Resources/UI/JNewPatient.ui')
        slicer.util.addWidget2(self, uiWidget)
        self.ui = slicer.util.childWidgetVariables(uiWidget)
        self.ui.btnClose.connect("clicked()", self.reject)
        self.ui.btn_create.connect("clicked()", self.on_save_clicked)
        self.ui.btn_save.connect("clicked()", self.on_save_clicked)
        self.ui.cmb_gender.clear()
        self.ui.cmb_degree.clear()
        self.ui.cmb_type.clear()
        self.ui.cmb_gender.addItems(self.gender_list)
        self.ui.cmb_degree.addItems(self.degree_list)
        self.ui.cmb_type.addItems(self.treat_type_list)
        self.current_date = datetime.now()
        self.ui.birthday.dateChanged.connect(self.set_age)
        self.ui.lbl_age.setEnabled(False)

    # info [birthday, name, age, id, note, gender, treatetype, degree, is_relapse]
    def set_dialog_type(self, is_create, info):
        self.patient_info = info
        if is_create:
            self.ui.btn_create.show()
            self.ui.btn_save.hide()
            self.ui.lbl_id.setText(info[3])
            return
        self.ui.label_12.setText("修改患者")
        self.ui.btn_create.hide()
        self.ui.btn_save.show()
        parsed_date = datetime.strptime(info[0], "%d %b %Y")
        date = qt.QDate(parsed_date.year, parsed_date.month, parsed_date.day)
        self.ui.birthday.setDate(date)
        self.ui.lbl_name.setText(info[1])
        self.ui.lbl_id.setText(info[3])
        self.ui.lbl_note.setText(info[4])
        gender = info[5]
        treat_type = info[6]
        degree_type = info[7]
        is_relapse = info[8]
        self.ui.cmb_gender.setCurrentIndex(gender)
        self.ui.cmb_type.setCurrentIndex(treat_type)
        self.ui.cmb_degree.setCurrentIndex(degree_type)
        if is_relapse == 1:
            self.ui.ckb_ulcer.setChecked(True)
        else:
            self.ui.ckb_ulcer.setChecked(False)

    def on_save_clicked(self):
        name = self.ui.lbl_name.text.strip()
        birthday = self.ui.birthday.text.strip()
        age = self.get_age(birthday)
        gender = self.ui.cmb_gender.currentIndex
        patient_id = self.ui.lbl_id.text.strip()
        treat_type = self.ui.cmb_type.currentIndex
        degree = self.ui.cmb_degree.currentIndex
        is_relapse = int(self.ui.ckb_ulcer.isChecked())
        note = self.ui.lbl_note.text.strip()
        print(name, age, gender, patient_id, treat_type, degree, is_relapse)
        if name == '' or patient_id == '':
            util.getModuleWidget("JMessageBox").show_infomation('提示', '名字或者患者ID为空, 保存失败', 0)
            return
        self.patient_info[0] = birthday
        self.patient_info[1] = name
        self.patient_info[2] = age
        self.patient_info[4] = note
        self.patient_info[5] = gender
        self.patient_info[6] = treat_type
        self.patient_info[7] = degree
        self.patient_info[8] = is_relapse
        self.accept()

    def set_age(self):
        self.ui.lbl_age.setText(str(self.get_age(self.ui.birthday.text)))

    def get_age(self, birthday):
        birth_date = datetime.strptime(birthday, "%d %b %Y")
        # 计算年龄
        age = self.current_date.year - birth_date.year
        # 如果今年还没过生日，年龄减一
        if (self.current_date.month, self.current_date.day) < (birth_date.month, birth_date.day):
            age -= 1
        return age


class ImageItem:
    ui = None
    widget = None
    item = None
    main = None

    def __init__(self, main, in_widget, in_ui) -> None:
        self.main = main
        self.widget = in_widget
        self.ui = in_ui

    def set_image(self, imageName):
        print(imageName)
        style_qss = f'image: url({imageName});'
        self.ui.lblImage.setStyleSheet(style_qss)


class VeinItem:
    ui = None
    widget = None
    item = None
    main = None
    segment_id = 0
    vein_name = ''

    def __init__(self, main, in_widget, in_ui) -> None:
        self.main = main
        self.widget = in_widget
        self.ui = in_ui
        self.ui.btn_detail.connect("clicked()", self.on_detail_click)

    def init_info(self, is_vein, info, segmentid, veinname):
        if is_vein == True:
            self.ui.btn_detail.hide()
        self.segment_id = segmentid
        self.vein_name = veinname + info
        self.ui.lblName.setText(info)

    def on_detail_click(self):
        self.main.show_details(self.vein_name, self.segment_id)


class JPCCS(ScriptedLoadableModule):

    def __init__(self, parent):
        ScriptedLoadableModule.__init__(self, parent)
        self.parent.title = "JPCCS"  # TODO: make this more human readable by adding spaces
        self.parent.categories = [
            "JPlugins"]  # TODO: set categories (folders where the module shows up in the module selector)
        self.parent.dependencies = []  # TODO: add here list of module names that this module requires
        self.parent.contributors = ["jia ze yu"]  # TODO: replace with "Firstname Lastname (Organization)"
        # TODO: update with short description of the module and a link to online module documentation
        self.parent.helpText = """

"""
        # TODO: replace with organization, grant and thanks
        self.parent.acknowledgementText = """

"""


class JPCCSWidget(ScriptedLoadableModuleWidget, VTKObservationMixin):
    browserWidget = None
    is_start = False
    timer = None
    init_once_flag = False
    patient_auto_id = -1
    conn = None
    cursor = None
    current_order = 0
    gender_list = ["男", "女", "其他"]
    # 治疗位置
    treat_type_list = ["左侧", "右侧"]
    # 临床分类
    degree_list = ["C1", "C2", "C3", "C4", "C5", "C6"]
    header_title_list = ['自增', '患者姓名', '患者生日', '患者年龄', '患者ID', '创建时间', '修改时间', '患者性别',
                         '治疗位置', '临床分类', '是否复发', '是否治疗', '备注', '操作']
    header_width_list = [10, 160, 160, 100, 240, 220, 220, 100, 140, 100, 100, 100, 120, 0]
    auto_id_column = 0
    name_column = 1
    birthday_column = 2
    age_column = 3
    id_column = 4
    create_column = 5
    modify_column = 6
    gender_column = 7
    treate_type_column = 8
    degree_column = 9
    relapse_column = 10
    treated_column = 11
    note_column = 12
    oeprate_column = 13

    def __init__(self, parent=None):
        """
        Called when the user opens the module the first time and the widget is initialized.
        """
        ScriptedLoadableModuleWidget.__init__(self, parent)
        VTKObservationMixin.__init__(self)  # needed for parameter node observation
        self.logic = None

    def setup(self):
        """
        Called when the user opens the module the first time and the widget is initialized.
        """
        ScriptedLoadableModuleWidget.setup(self)

        # Load widget from .ui file (created by Qt Designer).
        # Additional widgets can be instantiated manually and added to self.layout.
        self.logic = JPCCSLogic()
        self.logic.setWidget(self)

        self.uiWidget = slicer.util.loadUI(self.resourcePath('UI/JPCCS.ui'))
        self.layout.addWidget(self.uiWidget)
        self.ui = slicer.util.childWidgetVariables(self.uiWidget)

        self.uiWidget.setMRMLScene(slicer.mrmlScene)

        self.init_ui()
        self.init_extra_database()

    def init_extra_database(self):
        import sqlite3
        database_path = os.path.join(util.mainWindow().GetProjectBasePath(), "Resources", "Database", util.username,
                                     "ctkDICOM.sql")
        # 连接到数据库
        self.conn = sqlite3.connect(database_path)
        # 创建一个游标对象，用于执行 SQL 语句
        self.cursor = self.conn.cursor()

        self.cursor.execute('''
        CREATE TABLE IF NOT EXISTS Analysis (
            ID INTEGER PRIMARY KEY AUTOINCREMENT,
            PatientID INTEGER,
            pathlist TEXT,
            remark TEXT,
            recordtime DATETIME,
            modifytime DATETIME
        )
    ''')

        self.cursor.execute('''
        CREATE TABLE IF NOT EXISTS PatientInfo (
            ID INTEGER PRIMARY KEY AUTOINCREMENT,
            Name VARCHAR(100),
            Birthday VARCHAR(100),
            Age INTEGER,
            PatientID VARCHAR(52),
            CreateTime VARCHAR(100),
            ModifyTime VARCHAR(100),
            Gender INTEGER,
            TreatType INTEGER,
            Degree INTEGER,
            IsRelapse INTEGER,
            IsTreated INTEGER,
            Note VARCHAR(200),
            ImagePath VARCHAR(200)
        )
    ''')

        self.cursor.execute('''
        CREATE TABLE IF NOT EXISTS VeinInfo (
            ID INTEGER PRIMARY KEY AUTOINCREMENT,
            PatientID VARCHAR(52),
            VeinName VARCHAR(100)
        )
    ''')

        self.cursor.execute('''
        CREATE TABLE IF NOT EXISTS SegmentInfo (
            ID INTEGER PRIMARY KEY AUTOINCREMENT,
            VeinID INTEGER,
            SegmentName VARCHAR(100),
            IsTreated INTEGER,
            Pos VARCHAR(100),
            Power INTEGER,
            Energy INTEGER,
            ImageList TEXT
        )
    ''')

    # def init_once(self):
    #     if self.init_once_flag == True:
    #         return
    #     self.init_once_flag = True
    #     self.init_table()
    #     # menu = qt.QMenu("", self.ui.btn_export)
    #     # action_export_image = qt.QAction('导出影像', menu)
    #     # action_export_pdf = qt.QAction('导出报告', menu)
    #     # action_export_image.connect('triggered()', self.export_image)
    #     # action_export_pdf.connect('triggered()', self.export_pdf)
    #     # menu.addAction(action_export_image)
    #     # menu.addAction(action_export_pdf)
    #     # self.ui.btn_export.setMenu(menu)
    #     self.ui.btn_export.connect("clicked()", self.export_report)
    #
    #     menu2 = qt.QMenu("", self.ui.btn_shutdown)
    #     self.ui.btn_shutdown.setPopupMode(qt.QToolButton.InstantPopup)
    #     action_shutdown = qt.QAction('关机', menu2)
    #     action_sleep = qt.QAction('休眠', menu2)
    #     action_logout = qt.QAction('注销', menu2)
    #     action_shutdown.connect('triggered()', self.on_shutdown_clicked)
    #     action_sleep.connect('triggered()', self.on_sleep_clicked)
    #     action_logout.connect('triggered()', self.on_logout_clicked)
    #     menu2.addAction(action_shutdown)
    #     menu2.addAction(action_sleep)
    #     menu2.addAction(action_logout)
    #     self.ui.btn_shutdown.setMenu(menu2)

    def export_report(self):
        # 使用患者姓名和当前时间作为报告名
        patient_name = self.ui.table_patient.item(self.ui.table_patient.currentRow(), self.name_column).text()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(self.reports_path, f"{patient_name}_{timestamp}.docx")

        doc = Document()

        # 设置文档的默认字体
        def set_font(paragraph, font_name='宋体', font_size=12):
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # 添加标题
        heading = doc.add_heading('治疗报告', level=1)
        set_font(heading, font_name='宋体', font_size=20)

        # 添加患者基本信息
        doc.add_heading('基本信息', level=2)
        set_font(doc.paragraphs[-1], font_name='宋体', font_size=16)

        basic_info = {
            '姓名': patient_name,
            '性别': '男',
            '年龄': '35',
            '体检日期': '2024-07-08'
        }
        for key, value in basic_info.items():
            p = doc.add_paragraph(f'{key}: {value}')
            set_font(p, font_name='宋体', font_size=12)

        # 添加分隔线
        doc.add_paragraph('')

        # 添加检查结果
        doc.add_heading('检查结果', level=2)
        set_font(doc.paragraphs[-1], font_name='宋体', font_size=16)

        data = [
            ('项目', '结果'),
            ('血压', '120/80 mmHg'),
            ('血糖', '5.6 mmol/L'),
            ('胆固醇', '4.5 mmol/L'),
            ('心电图', '正常'),
            ('B超', '无异常')
        ]

        # 根据数据动态创建表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # 填充表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '项目'
        hdr_cells[1].text = '结果'
        set_font(hdr_cells[0].paragraphs[0], font_name='宋体', font_size=12)
        set_font(hdr_cells[1].paragraphs[0], font_name='宋体', font_size=12)

        # 填充表格数据
        for item, result in data[1:]:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = result
            set_font(row_cells[0].paragraphs[0], font_name='宋体', font_size=12)
            set_font(row_cells[1].paragraphs[0], font_name='宋体', font_size=12)

        # 添加分隔线
        doc.add_paragraph('')

        # 添加医生建议
        doc.add_heading('医生建议', level=2)
        set_font(doc.paragraphs[-1], font_name='宋体', font_size=16)

        advice = doc.add_paragraph(
            '根据上述检查结果，建议保持健康的生活方式，定期进行体检。如果有任何不适症状，请及时就医。')
        set_font(advice, font_name='宋体', font_size=12)

        # 添加图片
        doc.add_heading('附加图片', level=2)
        set_font(doc.paragraphs[-1], font_name='宋体', font_size=16)

        doc.add_paragraph('以下是检查报告相关的图片：')
        set_font(doc.paragraphs[-1], font_name='宋体', font_size=12)
        # doc.add_picture('1.jpg', width=Inches(4.0))

        # 保存文档
        doc.save(file_path)

        util.getModuleWidget("JMessageBox").show_infomation('提示', '导出报告成功', 0)

    def init_table(self):
        title_size = len(self.header_title_list)
        self.ui.table_patient.setColumnCount(title_size)
        self.ui.table_patient.setHorizontalHeaderLabels(self.header_title_list)
        self.ui.table_patient.setSelectionMode(qt.QAbstractItemView.SingleSelection)
        self.ui.table_patient.horizontalHeader().setFixedHeight(56)  # 固定高度
        self.ui.table_patient.verticalHeader().setDefaultSectionSize(56)
        self.ui.table_patient.setEditTriggers(qt.QAbstractItemView.NoEditTriggers)  # 设置不触发编辑选项
        self.ui.table_patient.horizontalHeader().setStretchLastSection(True)
        # 设置显示排序标志
        self.ui.table_patient.horizontalHeader().setSortIndicatorShown(True)
        # 隐藏纵向头部栏
        self.ui.table_patient.verticalHeader().hide()
        # 设置平分宽度,这里我们不需要平分
        # self.ui.table_patient.horizontalHeader().setSectionResizeMode(qt.QHeaderView.Stretch)
        for i in range(len(self.header_width_list)):
            self.ui.table_patient.setColumnWidth(i, self.header_width_list[i])

        self.ui.birthday.setLocale(qt.QLocale(qt.QLocale.Chinese))
        self.ui.birthday.setDisplayFormat("yyyy/MM/dd")

        self.ui.table_patient.setColumnHidden(self.auto_id_column, True)
        self.ui.table_patient.setColumnHidden(self.treate_type_column, True)
        self.ui.table_patient.setColumnHidden(self.degree_column, True)
        self.ui.table_patient.setColumnHidden(self.relapse_column, True)
        self.ui.table_patient.setColumnHidden(self.note_column, True)
        self.fetch_all_patient()

    def enter(self):
        self.init_table()
        self.addEvent(True)

    def exit(self):
        self.addEvent(False)

    def addEvent(self, bool_val):
        self.add_button_event("btn_delete", 'clicked()', self.on_delete_clicked, bool_val)
        self.add_button_event("btn_modify", 'clicked()', self.on_modify_clicked, bool_val)
        self.add_button_event("btn_new", 'clicked()', self.on_new_clicked, bool_val)
        self.add_button_event("btn_return", 'clicked()', self.on_return_clicked, bool_val)
        self.add_button_event("btn_search", 'clicked()', self.on_search_clicked, bool_val)
        self.add_button_event("btn_shutdown", 'clicked()', self.on_shutdown_clicked, bool_val)
        self.add_button_event("btnCopy", 'clicked()', self.on_copy_info, bool_val)
        self.add_button_event("btn_back", 'clicked()', self.back_to_manager, bool_val)
        self.add_button_event("btn_back2", 'clicked()', self.back_to_histroy, bool_val)
        self.add_button_event("cmbSelectType", 'currentIndexChanged(int)', self.on_select_type_change, bool_val)
        self.add_button_event("btn_export", "clicked()", self.export_report, bool_val)
        if bool_val:
            self.ui.table_patient.horizontalHeader().connect('sectionClicked(int)', self.on_sort)
            self.ui.table_patient.connect('currentCellChanged(int, int, int, int)', self.on_table_select_change)
            # self.ui.table_patient.connect('currentItemChanged(QListWidgetItem*,QListWidgetItem*)', self.on_table_select_change)
        else:
            self.ui.table_patient.horizontalHeader().disconnect('sectionClicked(int)', self.on_sort)
            self.ui.table_patient.disconnect('currentCellChanged(int, int, int, int)', self.on_table_select_change)
            # self.ui.table_patient.disconnect('currentItemChanged(QListWidgetItem*,QListWidgetItem*)', self.on_table_select_change)

    def on_table_select_change(self, currRow, currColumn, perviouRow, perviousColumn):
        # print("currRow, currColumn, perviouRow, perviousColumn", currRow, currColumn, perviouRow, perviousColumn)
        detail_info, details = self.get_user_detail_info(currRow)
        self.ui.lblInfo.setPlainText(detail_info)
        self.ui.btn_modify.setEnabled(True)
        self.ui.btn_export.setEnabled(False)
        if details['状态'] == "已治疗":
            self.ui.btn_modify.setEnabled(False)
            self.ui.btn_export.setEnabled(True)
        self.current_select_row = currRow

    def on_copy_info(self):
        util.add_to_clipboard(self.ui.lblInfo.text)

    def back_to_manager(self):
        self.ui.tabWidget.setCurrentIndex(0)

    def back_to_histroy(self):
        self.ui.tabWidget.setCurrentIndex(2)

    def on_select_type_change(self, select_type):
        index = self.ui.cmbSelectType.currentIndex
        if index == 0:
            self.fetch_all_patient()
        else:
            self.fetch_patient_by_type(select_type - 1)
        pass

    def get_user_detail_info(self, row):
        columns = {
            "患者姓名": self.name_column,
            "患者生日": self.birthday_column,
            "患者年龄": self.age_column,
            "患者ID": self.id_column,
            "创建时间": self.create_column,
            "修改时间": self.modify_column,
            "患者性别": self.gender_column,
            "治疗位置": self.treate_type_column,
            "临床分类": self.degree_column,
            "是否复发": self.relapse_column,
            "状态": self.treated_column,
            "备注": self.note_column
        }

        details = {}
        for key, column in columns.items():
            details[key] = self.ui.table_patient.item(row, column).text()

        detail_info = "\n".join([f"{key}:  {value}" for key, value in details.items()])

        return detail_info, details
        user_name = self.ui.table_patient.item(row, self.name_column).text()
        user_birthday = self.ui.table_patient.item(row, self.birthday_column).text()
        user_age = self.ui.table_patient.item(row, self.age_column).text()
        user_id = self.ui.table_patient.item(row, self.id_column).text()
        create_date = self.ui.table_patient.item(row, self.create_column).text()
        modify_date = self.ui.table_patient.item(row, self.modify_column).text()
        user_gender = self.ui.table_patient.item(row, self.gender_column).text()
        treate_pos = self.ui.table_patient.item(row, self.treate_type_column).text()
        kind = self.ui.table_patient.item(row, self.degree_column).text()
        relapse = self.ui.table_patient.item(row, self.relapse_column).text()
        note = self.ui.table_patient.item(row, self.note_column).text()
        treated = self.ui.table_patient.item(row, self.treated_column).text()
        detail_info = f'''
    患者姓名:  {user_name}
    患者生日:  {user_birthday}
    患者年龄:  {user_age}
    患者ID:  {user_id}
    创建时间:  {create_date}
    修改时间:  {modify_date}
    患者性别:  {user_gender}
    治疗位置:  {treate_pos}
    临床分类:  {kind}
    是否复发:  {relapse}
    状态:  {treated}
    备注:  {note}'''
        return detail_info

    def on_sort(self, column):
        if self.current_order == qt.Qt.AscendingOrder:
            self.current_order = qt.Qt.DescendingOrder
        else:
            self.current_order = qt.Qt.AscendingOrder
        self.ui.table_patient.horizontalHeader().setSortIndicator(column, self.current_order)
        self.ui.table_patient.sortByColumn(column)

    def on_delete_clicked(self):
        current_row = self.ui.table_patient.currentRow()
        if current_row == -1:
            util.messageBox("请先选择一行进行删除")
            return
        delete_id = int(self.ui.table_patient.item(current_row, 0).text())
        delete_name = self.ui.table_patient.item(current_row, 1).text()
        result = util.getModuleWidget("JMessageBox").show_question("提示", f'是否要删除姓名为{delete_name}患者？')
        if result == 0:
            return
        self.cursor.execute('DELETE from PatientInfo where ID=?', (delete_id,))
        self.conn.commit()
        self.fetch_all_patient()

    def on_new_clicked(self):
        unique_id = util.get_unique_id()
        patient_info = ['', '', '', unique_id, '', 0, 0, 0, 0]
        dialog = JNewPatient(slicer.util.mainWindow())
        # print(f"patient info {patient_info}")
        dialog.set_dialog_type(True, patient_info)
        result = dialog.exec_()
        if result == 0:
            return
        # [birthday, name, age, id, note, gender, treatetype, degree, is_relapse]
        registration_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.cursor.execute('''INSERT INTO PatientInfo (Name, Age, Birthday, PatientID, CreateTime, ModifyTime, Gender, TreatType, Degree, IsRelapse, Note, IsTreated)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''', (
            patient_info[1], patient_info[2], patient_info[0], unique_id, registration_time, registration_time,
            patient_info[5], patient_info[6], patient_info[7], patient_info[8], patient_info[4], 0))
        try:
            self.conn.commit()
            self.fetch_all_patient()
        except TypeError:
            pass

    def on_next_clicked(self):
        util.send_event_str(util.GotoNextPage)
        pass

    def on_return_clicked(self):
        self.ui.tabWidget.setCurrentIndex(0)
        self.fetch_all_patient()

    def fetch_all_patient(self):
        import sqlite3
        database_path = os.path.join(util.mainWindow().GetProjectBasePath(), "Resources", "Database", util.username,
                                     "ctkDICOM.sql")
        self.cursor.execute(f'SELECT * FROM PatientInfo')
        items_list = self.cursor.fetchall()
        self.set_table_data(items_list)

    def fetch_patient_by_type(self, select_type):
        self.cursor.execute(f'SELECT * FROM PatientInfo WHERE IsTreated = ?', (select_type,))
        items_list = self.cursor.fetchall()
        self.set_table_data(items_list)

    def set_table_data(self, datas):
        self.ui.table_patient.setRowCount(len(datas))
        row_index = 0
        for item in datas:
            id = item[self.auto_id_column]
            name = item[self.name_column]
            for i in range(len(self.header_title_list)):
                info = self.get_des_info(i, item[i])
                widget_item = qt.QTableWidgetItem(info)
                widget_item.setData(qt.Qt.UserRole, item[i])
                widget_item.setTextAlignment(qt.Qt.AlignCenter)
                self.ui.table_patient.setItem(row_index, i, widget_item)
            is_treated = item[self.treated_column]
            btn_info = "配置"
            if is_treated == 1:
                btn_info = "记录"
            btn_detail = qt.QPushButton(btn_info)
            btn_detail.setFixedSize(100, 35)

            # 实现居中
            widget = qt.QWidget()
            layout = qt.QHBoxLayout()
            layout.setSpacing(0)
            layout.setMargin(0)
            layout.addWidget(btn_detail)
            widget.setLayout(layout)

            self.ui.table_patient.setCellWidget(row_index, self.oeprate_column, widget)

            btn_detail.clicked.connect(
                lambda obj, id=id, name=name, is_treated=is_treated: self.on_detail_click(id, name, is_treated))
            row_index = row_index + 1

    def on_detail_click(self, id, name, is_treated):
        if is_treated == False:
            util.global_data_map["patientID"] = id
            util.global_data_map["patientName"] = name
            util.global_data_map['UltraSoundSettingPanel'].on_reset_ui()
            util.send_event_str(util.SetPage, 4)
            return
        self.ui.tabWidget.setCurrentIndex(2)
        self.cursor.execute(f'SELECT * FROM VeinInfo WHERE PatientID = ?', (id,))
        items_list = self.cursor.fetchall()
        self.ui.listWidget.clear()
        print(items_list)
        for item in items_list:
            vein_name = item[2]
            self.create_vein_item(True, vein_name, 0, vein_name)
            index = 0
            vein_id = item[0]
            self.cursor.execute(f'SELECT * FROM SegmentInfo WHERE VeinID = ?', (vein_id,))
            segment_list = self.cursor.fetchall()
            for segment in segment_list:
                self.create_vein_item(False, segment[2], segment[0], vein_name)
                index = index + 1
        pass

    def show_segment_detail(self, des, segment_id):
        self.ui.tabWidget.setCurrentIndex(3)
        self.ui.lbl_segment_des.setText(des)
        self.ui.imageList.clear()
        self.cursor.execute(f'SELECT * FROM SegmentInfo WHERE ID = ?', (segment_id,))
        item = self.cursor.fetchone()
        detail_info = self.get_segment_des_info(item)
        self.ui.lbl_segment_info.setText(detail_info)
        images = item[6].split(";")
        for image_name in images:
            image_name = image_name.replace("\\", "/")
            self.create_image_item(image_name)
        pass

    def show_details(self, des, segment_id):
        self.cursor.execute("SELECT Pos,Power,Energy,ImageList FROM SegmentInfo WHERE ID = ?", (segment_id,))
        detail_info = self.cursor.fetchone()
        history_detail_panel = HistoryDetailPanel(des, detail_info)
        history_detail_panel.show()

    def create_image_item(self, imageName):
        template1 = slicer.util.loadUI(self.resourcePath("UI/ImageItem.ui"))
        template1ui = slicer.util.childWidgetVariables(template1)
        template = ImageItem(self, template1, template1ui)
        item = qt.QListWidgetItem(self.ui.imageList)
        template.set_image(imageName)
        height_of_list = self.ui.imageList.height
        height_of_item = height_of_list - 250
        width_of_item = height_of_item * 0.8
        item.setSizeHint(qt.QSize(width_of_item, height_of_item))
        self.ui.imageList.setItemWidget(item, template.widget)
        self.ui.imageList.addItem(item)

    def create_vein_item(self, is_vein, text, segmentId, veinname):
        template1 = slicer.util.loadUI(self.resourcePath("UI/VeinItem.ui"))
        template1ui = slicer.util.childWidgetVariables(template1)
        template = VeinItem(self, template1, template1ui)
        item = qt.QListWidgetItem(self.ui.listWidget)
        template.init_info(is_vein, text, segmentId, veinname)
        item.setSizeHint(qt.QSize(340, 50))
        self.ui.listWidget.setItemWidget(item, template.widget)
        self.ui.listWidget.addItem(item)

    def get_segment_des_info(self, item):
        detail_info = f'''
    ID:  {item[0]}
    片段名称:  {item[2]}
    位置:  {item[3]}
    治疗功率:  {item[4]}
    治疗能量:  {item[5]}'''
        return detail_info

    def get_des_info(self, index, info):
        result = str(info)
        print(index, result)
        if index == self.gender_column:
            result = self.gender_list[info]
        elif index == self.treate_type_column:
            result = self.treat_type_list[info]
        elif index == self.degree_column:
            result = self.degree_list[info]
        elif index == self.relapse_column:
            result = "否"
            if info == 1:
                result = "是"
        elif index == self.treated_column:
            result = "未治疗"
            if info == 1:
                result = "已治疗"
        return result

    def on_search_clicked(self):
        search_str = self.ui.lbl_search_info.text.strip()
        if search_str == '':
            self.fetch_all_patient()
            return
        search_type = self.ui.cmb_search_type.currentIndex
        find_type = 'Name'
        if search_str == 1:
            find_type = 'PatientID'
        search_info = f'%{search_str}%'
        self.cursor.execute('SELECT * FROM PatientInfo WHERE {} LIKE ?'.format(find_type), (search_info,))
        user_data = self.cursor.fetchall()
        self.set_table_data(user_data)
        pass

    def on_shutdown_clicked(self):
        util.shut_down()
        pass

    def on_sleep_clicked(self):
        util.shut_sleep()
        pass

    def on_logout_clicked(self):
        util.shut_logout()
        pass

    def on_modify_clicked(self):
        current_row = self.ui.table_patient.currentRow()
        if current_row == -1:
            util.messageBox("请先选择一行修改")
            return
        # '自增',  '',   '患者姓名', '患者生日', '患者年龄', '患者ID', '患者性别', '治疗位置', '临床分类', '是否复发', '备注'

        patient_info = ['', '', '', '', '', 0, 0, 0,
                        0]  # [birthday, name, age, id, note, gender, treatetype, degree, is_relapse]
        patient_auto_id = int(self.ui.table_patient.item(current_row, self.auto_id_column).text())
        patient_info[0] = self.ui.table_patient.item(current_row, self.birthday_column).text()
        patient_info[1] = self.ui.table_patient.item(current_row, self.name_column).text()
        patient_info[2] = self.ui.table_patient.item(current_row, self.age_column).text()
        patient_info[3] = self.ui.table_patient.item(current_row, self.id_column).text()
        patient_info[4] = self.ui.table_patient.item(current_row, self.note_column).text()
        patient_info[5] = int(self.ui.table_patient.item(current_row, self.gender_column).data(qt.Qt.UserRole))
        patient_info[6] = int(self.ui.table_patient.item(current_row, self.treate_type_column).data(qt.Qt.UserRole))
        patient_info[7] = int(self.ui.table_patient.item(current_row, self.degree_column).data(qt.Qt.UserRole))
        patient_info[8] = int(self.ui.table_patient.item(current_row, self.relapse_column).data(qt.Qt.UserRole))

        dialog = JNewPatient(slicer.util.mainWindow())
        dialog.set_dialog_type(False, patient_info)
        result = dialog.exec_()
        if result == 0:
            return
        # [birthday, name, age, id, note, gender, treatetype, degree, is_relapse]
        modify_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.cursor.execute('''UPDATE PatientInfo SET Name=?, Age=?, Birthday=?, PatientID=?, ModifyTime=?, Gender=?, TreatType=?, Degree=?, IsRelapse=?, Note=?
        WHERE ID=?''', (
            patient_info[1], patient_info[2], patient_info[0], patient_info[3], modify_time, patient_info[5],
            patient_info[6],
            patient_info[7], patient_info[8], patient_info[4], patient_auto_id))
        try:
            self.conn.commit()
            self.fetch_all_patient()
            detail_info, details = self.get_user_detail_info(current_row)
            self.ui.lblInfo.setPlainText(detail_info)
        except TypeError:
            pass
        pass

    def add_button_event(self, btn_name, operate_type, callback_func, is_connect=True):
        if hasattr(self.ui, btn_name) == False:
            return
        button = getattr(self.ui, btn_name)
        if is_connect == True:
            button.connect(operate_type, callback_func)
        else:
            button.disconnect(operate_type, callback_func)

    def refresh_patient_analysis(self):
        return

    def on_trigger(self):
        pass

    '''
      当有新的ScalarVolumeNode添加的时候,恢复初始设置
    '''

    @vtk.calldata_type(vtk.VTK_OBJECT)
    def OnMainNodeAdded(self, caller, str_event, calldata):
        self.logic.m_Node = calldata
        if self.logic.m_Node:
            util.SetGlobalSaveValue("JPCCS_MainNodeID", self.logic.m_Node.GetID())

    '''
      当原有的ScalarVolumeNode删除的时候,删除所有的标签
    '''

    @vtk.calldata_type(vtk.VTK_OBJECT)
    def OnMainNodeRemoved(self, caller, str_event, calldata):
        self.logic.m_Node = None
        util.SetGlobalSaveValue("JPCCS_MainNodeID", None)

    def OnArchiveLoaded(self, _a, _b):
        print(f"load archive from {self.__class__.__name__}")
        nodeid = util.GetGlobalSaveValue("JPCCS_MainNodeID")
        print("OnArchiveLoaded JPCCS main nodeid is", nodeid)
        if nodeid is None:
            return
        node = util.GetNodeByID(nodeid)
        self.logic.m_Node = node

    def init_later(self):
        pass

    def init_ui(self):
        self.ui.lbl_age.setEnabled(False)
        self.ui.tabWidget.tabBar().hide()
        util.singleShot(10, self.init_later)
        self.ui.cmb_gender.clear()
        self.ui.cmb_degree.clear()
        self.ui.cmb_type.clear()
        self.ui.cmb_gender.addItems(self.gender_list)
        self.ui.cmb_degree.addItems(self.degree_list)
        self.ui.cmb_type.addItems(self.treat_type_list)
        self.ui.btnCopy.setVisible(False)
        self.ui.lblInfo.setReadOnly(True)
        self.ui.lbl_search_info.setPlaceholderText('请输入搜索关键字')

        self.reports_path = os.path.join(util.get_project_base_path(), "Resources", "reports").replace("\\", "/")
        if not os.path.exists(self.reports_path):
            os.makedirs(self.reports_path)


class JPCCSLogic(ScriptedLoadableModuleLogic):
    m_Node = None

    def __init__(self):
        """
        Called when the logic class is instantiated. Can be used for initializing member variables.
        """
        ScriptedLoadableModuleLogic.__init__(self)
        slicer.mrmlScene.AddObserver(slicer.vtkMRMLScene.NodeAddedEvent, self.onNodeAdded)

    def setWidget(self, widget):
        self.m_Widget = widget

    @vtk.calldata_type(vtk.VTK_OBJECT)
    def onNodeAdded(self, caller, event, calldata):
        node = calldata
        # if isinstance(node, slicer.vtkMRMLMarkupsFiducialNode):
        # self.m_Widget.on_node_added(node)
